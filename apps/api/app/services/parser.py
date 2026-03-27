from __future__ import annotations

"""
Resume parser service.

Extracts plain text from PDF and DOCX files. Structured field extraction
(name, email, skills, experience) is a separate concern handled by the
scoring / LLM layer.

Current implementation: placeholder returning empty strings.

TODO:
  - PDF: use ``pdfminer.six`` or ``pypdf`` for text extraction.
  - DOCX: use ``python-docx`` to walk paragraphs and tables.
  - Fallback: OCR via Azure Document Intelligence for scanned PDFs.
"""

import io
import logging
from enum import StrEnum

logger = logging.getLogger(__name__)


class DocumentFormat(StrEnum):
    PDF = "application/pdf"
    DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    DOC = "application/msword"
    UNKNOWN = "application/octet-stream"


class ParserError(Exception):
    """Raised when text extraction fails."""


async def extract_text(data: bytes, content_type: str) -> str:
    """
    Extract plain text from a resume file.

    Args:
        data: Raw bytes of the uploaded file.
        content_type: MIME type of the file, used to select the parser.

    Returns:
        Extracted plain text (may be empty if no text layer is found).

    Raises:
        ParserError: If extraction fails catastrophically.
    """
    fmt = DocumentFormat(content_type) if content_type in DocumentFormat._value2member_map_ else DocumentFormat.UNKNOWN  # type: ignore[attr-defined]

    logger.info("Extracting text from %s document (%d bytes)", fmt, len(data))

    if fmt == DocumentFormat.PDF:
        return await _extract_from_pdf(data)
    if fmt in (DocumentFormat.DOCX, DocumentFormat.DOC):
        return await _extract_from_docx(data)

    raise ParserError(f"Unsupported document format: {content_type}")


async def _extract_from_pdf(data: bytes) -> str:
    """
    Extract text from a PDF using pypdf.

    Iterates over every page and joins extracted text with newlines.
    Falls back to Azure Document Intelligence for scanned / image-based PDFs
    that contain no text layer (future TODO).
    """
    try:
        import pypdf  # noqa: PLC0415 — lazy import keeps startup fast when unused

        reader = pypdf.PdfReader(io.BytesIO(data))
        pages: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)

        text = "\n".join(pages)
        logger.info("PDF extraction complete: %d page(s), %d characters", len(pages), len(text))
        return text
    except Exception as exc:
        raise ParserError(f"PDF extraction failed: {exc}") from exc


async def _extract_from_docx(data: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.

    Collects text from all paragraphs and from every cell in every table,
    joining everything with newlines to preserve visual structure.
    """
    try:
        import docx  # noqa: PLC0415 — lazy import keeps startup fast when unused

        doc = docx.Document(io.BytesIO(data))

        parts: list[str] = []

        # Body paragraphs (headings, normal text, bullet points, etc.)
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Tables — iterate rows and cells so tabular data is not lost
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append("\t".join(row_cells))

        text = "\n".join(parts)
        logger.info(
            "DOCX extraction complete: %d paragraph(s), %d table(s), %d characters",
            len(doc.paragraphs),
            len(doc.tables),
            len(text),
        )
        return text
    except Exception as exc:
        raise ParserError(f"DOCX extraction failed: {exc}") from exc


async def chunk_text(text: str, chunk_size: int = 512, overlap: int = 64) -> list[str]:
    """
    Split a long text into overlapping chunks suitable for embedding.

    Args:
        text: Full plain-text resume content.
        chunk_size: Target number of tokens/characters per chunk.
        overlap: Number of characters to overlap between adjacent chunks.

    Returns:
        List of text chunks.

    TODO: Replace character-based chunking with a proper tokenizer-aware
          splitter (e.g. ``tiktoken`` with ``cl100k_base``).
    """
    if not text:
        return []

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        start += chunk_size - overlap

    return chunks
